import os
import glob
import threading
import customtkinter as ctk
import re  # 添加re模块用于正则表达式
from tkinter import filedialog, messagebox
from docx import Document
from docxcompose.composer import Composer
import win32com.client as win32
from time import sleep

ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")


class WordMergerApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Word文档合并工具")
        self.geometry("800x600")
        self.selected_dir = ""
        self.merge_algorithm = "simple"  # 默认合并算法
        self.create_widgets()
        self.file_page_map = {}  # 添加文件页码映射字典
        # 检查是否是Windows系统，如果是，显示提示消息
        if os.name == 'nt':
            messagebox.showwarning(
                "Windows系统提示", 
                "你使用的是Windows系统，请保存和关闭所有打开的Word文档，以免造成文档丢失。关闭后再运行合并程序。"
            )
    def create_widgets(self):
        # 目录选择部分
        self.dir_frame = ctk.CTkFrame(self)
        self.dir_frame.pack(pady=10, padx=10, fill="x")

        self.dir_button = ctk.CTkButton(self.dir_frame,text="选择目录",command=self.select_directory)
        self.dir_button.pack(side="left", padx=5)

        self.dir_label = ctk.CTkLabel(self.dir_frame,text="未选择目录",text_color='black',anchor="w")
        self.dir_label.pack(side="left", padx=5)
        # 合并算法选择部分
        self.algorithm_frame = ctk.CTkFrame(self)
        self.algorithm_frame.pack(pady=10, padx=10, fill="x")

        self.algorithm_label = ctk.CTkLabel(self.algorithm_frame,text="选择合并算法：",anchor="w")
        self.algorithm_label.pack(side="left", padx=5)

        # 合并算法选项
        self.algorithm_var = ctk.StringVar(value="simple")
        self.algorithm_simple = ctk.CTkRadioButton(self.algorithm_frame,text="简单追加",variable=self.algorithm_var,value="simple")
        self.algorithm_simple.pack(side="left", padx=5)

        self.algorithm_format = ctk.CTkRadioButton(self.algorithm_frame,text="保留格式",variable=self.algorithm_var,value="format")
        self.algorithm_format.pack(side="left", padx=5)

        self.algorithm_word_api = ctk.CTkRadioButton(self.algorithm_frame,text="使用 Word API",variable=self.algorithm_var,value="word_api")
        self.algorithm_word_api.pack(side="left", padx=5)

        self.algorithm_docxcompose = ctk.CTkRadioButton(self.algorithm_frame,text="使用 docxcompose",variable=self.algorithm_var,value="docxcompose")
        self.algorithm_docxcompose.pack(side="left", padx=5)

        # 日志显示部分
        self.log_text = ctk.CTkTextbox(self, wrap="none")
        self.log_text.pack(pady=10, padx=10, fill="both", expand=True)

        # 合并按钮
        self.merge_button = ctk.CTkButton(self,text="开始合并",command=self.start_merge,state="disabled")
        self.merge_button.pack(pady=10)

    def select_directory(self):
        """选择目录"""
        self.selected_dir = filedialog.askdirectory()
        if self.selected_dir:
            self.dir_label.configure(text=self.selected_dir)
            self.merge_button.configure(state="normal")
            self.log("已选择目录：" + self.selected_dir)

    def log(self, message):
        """在日志框中显示消息"""
        self.log_text.configure(state="normal")
        self.log_text.insert("end", message + "\n")
        self.log_text.see("end")
        self.log_text.configure(state="disabled")

    def start_merge(self):
        """启动合并线程"""
        self.merge_algorithm = self.algorithm_var.get()  # 获取选择的合并算法
        threading.Thread(target=self.merge_documents, daemon=True).start()

    def merge_documents(self):
        """合并文档主逻辑"""
        try:
            # 检查目录有效性
            if not os.path.isdir(self.selected_dir):
                self.log("错误：目录不存在")
                messagebox.showerror("目录错误", "选择的目录不存在或已被删除")
                return

            # 获取所有Word文档（过滤掉以~$开头的缓存文件）
            doc_files = glob.glob(os.path.join(self.selected_dir, "*.doc*"))
            doc_files = [
                f for f in doc_files
                if f.endswith((".doc", ".docx")) and not os.path.basename(f).startswith("~$")
            ]
            doc_files.sort()

            if not doc_files:
                self.log("错误：目录中没有找到Word文档")
                messagebox.showerror("文件未找到", "目录中没有有效的Word文档")
                return

            # 设置输出路径
            output_dir = os.path.join(self.selected_dir, "合并结果")
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, "合并完成文档.docx")

            # 根据选择的合并算法执行合并
            if self.merge_algorithm == "simple":
                if os.name=='nt':
                    success = self.algorithm_windows(doc_files, output_path)
                else:
                    success = self.merge_simple(output_path, doc_files)
            elif self.merge_algorithm == "format":
                success = self.merge_with_format(output_path, doc_files)
            elif self.merge_algorithm == "word_api":
                success = self.merge_with_word_api(output_path, doc_files)
            elif self.merge_algorithm == "docxcompose":
                success = self.merge_with_docxcompose(output_path, doc_files)
            else:
                self.log("错误：未知的合并算法")
                return

            if success:
                # 生成目录
                self.generate_toc(output_path)
                self.log("\n合并完成！文件已保存到：" + output_path)
                messagebox.showinfo("完成", "文档合并完成！")
            else:
                self.log("合并失败，请检查日志")
                messagebox.showerror("错误", "合并失败，请检查日志")

        except Exception as e:
            error_msg = f"合并过程中发生严重错误：{str(e)}"
            self.log(error_msg)
            messagebox.showerror("严重错误", error_msg)

    def extract_display_name(self, filename):
        """提取带书名号的显示名称，没有书名号则用原文件名（不含扩展名）"""
        # 去除文件扩展名
        name_without_ext = os.path.splitext(filename)[0]
        
        # 清理文件名中可能的特殊字符
        name_without_ext = name_without_ext.strip()
        
        # 查找书名号内容
        match = re.search(r"《(.+?)》", name_without_ext)
        
        if match:
            # 如果有书名号则取内容
            display_name = match.group(1).strip()
            # 确保提取的内容不为空
            if not display_name:
                display_name = name_without_ext
        else:
            # 否则保留整个文件名（不含扩展名）
            display_name = name_without_ext
        
        # 限制长度，防止过长的文件名
        if len(display_name) > 100:
            display_name = display_name[:97] + "..."
            
        # 移除可能导致问题的特殊字符
        display_name = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', display_name)
        
        return display_name

    def generate_toc(self, doc_path):
        """生成目录"""
        word = None
        try:
            self.log("正在生成目录...")
            
            # 检查文件路径是否存在
            if not os.path.exists(doc_path):
                raise FileNotFoundError(f"文件未找到：{doc_path}")
            
            # 强制关闭所有Word进程
            self.log("尝试关闭所有Word进程...")
            try:
                import subprocess
                subprocess.call("taskkill /f /im WINWORD.EXE", shell=True)
                self.log("已强制关闭所有Word进程")
                # 等待进程完全关闭
                sleep(3)
            except Exception as e:
                self.log(f"关闭Word进程时出错: {str(e)}")
            
            # 使用直接修改原文档的方法
            self.log("在原文档中生成目录...")
            word = win32.Dispatch("Word.Application")
            word.Visible = False
            
            # 打开原文档
            self.log(f"打开文档: {doc_path}")
            doc = word.Documents.Open(doc_path)
            
            # 在文档开头插入目录
            # 不使用分页符，直接在开头插入
            doc.Range(0, 0).InsertParagraphBefore()
            
            # 插入目录标题
            title_range = doc.Range(0, 0)
            title_range.InsertBefore("目录\r\n\r\n")
            title_range = doc.Range(0, len("目录\r\n"))
            title_range.Font.Name = "宋体"
            title_range.Font.Size = 16  # 三号字体约为16磅
            title_range.Font.Bold = True
            title_range.ParagraphFormat.Alignment = 1  # 居中
            
            # 设置单倍行距
            title_range.ParagraphFormat.LineSpacing = 12  # 单倍行距
            title_range.ParagraphFormat.LineSpacingRule = 0  # 0 = wdLineSpaceSingle
            
            # 添加空行
            doc.Range(len("目录\r\n\r\n"), len("目录\r\n\r\n")).InsertParagraphAfter()
            
            # 获取文件列表并按顺序排序
            file_list = list(self.file_page_map.keys())
            file_list.sort()
            
            # 当前插入位置 - 在目录标题之后
            current_pos = title_range.End + 1
            
            # 添加目录项
            for i, file_path in enumerate(file_list):
                try:
                    # 使用extract_display_name方法提取显示名称
                    display_name = self.extract_display_name(os.path.basename(file_path))
                    
                    # 限制显示名称长度，防止过长
                    if len(display_name) > 100:
                        display_name = display_name[:97] + "..."
                    
                    # 使用文件页码映射中的实际页码
                    page_number = self.file_page_map[file_path]['page'] + 1  # +1 因为目录页
                    bookmark = self.file_page_map[file_path]['bookmark']
                    
                    self.log(f"添加目录项: {display_name}, 页码: {page_number}")
                    
                    # 创建完整的目录行
                    complete_line = f"{display_name}\t{page_number}"
                    
                    # 插入目录项
                    current_range = doc.Range(current_pos, current_pos)
                    current_range.Text = complete_line
                    
                    # 设置段落格式
                    current_para = current_range.Paragraphs(1)
                    current_para.Range.Font.Name = "宋体"
                    current_para.Range.Font.Size = 16  # 三号字体约为16磅
                    
                    # 设置单倍行距
                    current_para.Format.LineSpacing = 12  # 单倍行距
                    current_para.Format.LineSpacingRule = 0  # 0 = wdLineSpaceSingle
                    
                    # 添加制表符和点线
                    current_para.Format.TabStops.ClearAll()
                    current_para.Format.TabStops.Add(Position=450, Alignment=2, Leader=1)  # 点线制表符
                    
                    # 创建超链接
                    try:
                        # 计算链接范围
                        link_start = current_para.Range.Start
                        link_end = current_para.Range.End - 1  # -1 避免包含段落标记
                        link_range = doc.Range(link_start, link_end)
                        
                        # 创建超链接
                        doc.Hyperlinks.Add(
                            Anchor=link_range,
                            SubAddress=bookmark
                        )
                        self.log(f"超链接创建成功: {display_name}")
                    except Exception as e:
                        self.log(f"创建超链接时出错: {str(e)}")
                    
                    # 添加换行符（除了最后一个条目）
                    if i < len(file_list) - 1:
                        current_para.Range.InsertParagraphAfter()
                    
                    # 更新当前位置
                    current_pos = doc.Range(0, current_para.Range.End).End
                    
                except Exception as e:
                    self.log(f"处理目录项时出错: {str(e)}")
                    continue
            
            # 添加单个分页符，将目录与正文分开
            self.log("在目录后添加分页符...")
            doc.Range(current_pos, current_pos).InsertBreak(7)  # 7 = wdPageBreak
            
            # 保存文档
            self.log("保存文档...")
            doc.Save()
            doc.Close(SaveChanges=True)
            
            # 关闭Word
            word.Quit()
            self.log("目录生成完成")
            return True
            
        except Exception as e:
            self.log(f"生成目录时出错：{str(e)}")
            import traceback
            self.log(traceback.format_exc())
            return False
        finally:
            # 确保在任何情况下都关闭Word
            try:
                if word:
                    word.Quit()
            except:
                pass
            
            # 最后一次尝试强制关闭所有Word进程
            try:
                import subprocess
                subprocess.call("taskkill /f /im WINWORD.EXE", shell=True)
            except:
                pass

    def merge_simple(self, output_path, doc_files):
        """简单追加合并算法，跨平台支持"""
        try:
            merged_doc = Document()
            self.file_page_map = {}  # 重置文件页码映射
            current_page = 0
            
            # 预处理：在Windows系统上将.doc文件转换为.docx
            temp_files = []
            is_windows = os.name == 'nt'
            
            for i, file_path in enumerate(doc_files):
                try:
                    self.log(f"正在处理文件：{os.path.basename(file_path)}")
                    file_path = os.path.abspath(file_path)  # 确保使用绝对路径
                    
                    # 如果是Windows系统且是.doc文件，先转换为.docx
                    if is_windows and file_path.lower().endswith('.doc'):
                        self.log(f"转换.doc文件为.docx: {os.path.basename(file_path)}")
                        # 使用Word COM接口转换
                        word = None
                        try:
                            word = win32.Dispatch("Word.Application")
                            word.Visible = False
                            
                            # 打开.doc文件
                            doc = word.Documents.Open(file_path)
                            # 创建临时.docx文件路径
                            temp_docx = os.path.join(os.path.dirname(output_path), f"temp_{i}_{os.path.basename(file_path)}x")
                            # 保存为.docx
                            doc.SaveAs(temp_docx, 16)  # 16 = wdFormatDocumentDefault (.docx)
                            doc.Close(SaveChanges=False)
                            
                            # 添加到临时文件列表
                            temp_files.append(temp_docx)
                            # 使用转换后的文件
                            file_path = temp_docx
                            self.log(f"成功转换文件：{os.path.basename(file_path)}")
                        except Exception as e:
                            self.log(f"转换文件失败: {str(e)}")
                            continue
                        finally:
                            if word:
                                try:
                                    word.Quit()
                                except:
                                    pass
                    # 非Windows系统下，跳过.doc文件
                    elif not is_windows and file_path.lower().endswith('.doc'):
                        self.log(f"跳过.doc文件（非Windows系统不支持）: {os.path.basename(file_path)}")
                        continue
                    
                    # 创建唯一书签名
                    bookmark_name = f"bookmark_{i+1}"
                    
                    try:
                        # 尝试打开文件
                        doc = Document(file_path)
                        
                        # 估算页数 (粗略计算，每页约2000个字符)
                        text_length = 0
                        for para in doc.paragraphs:
                            text_length += len(para.text)
                        # 更新当前页码 (粗略估计)
                        page_count = max(1, text_length // 2000)
                        
                        # 记录当前页码和书签
                        self.file_page_map[file_path] = {
                            'page': current_page,
                            'bookmark': bookmark_name
                        }
                        
                        # 追加内容
                        for paragraph in doc.paragraphs:
                            merged_doc.add_paragraph(paragraph.text)
                        
                        # 添加分页符（除了最后一个文档）
                        if i < len(doc_files) - 1:
                            merged_doc.add_page_break()
                        
                        current_page += page_count
                        self.log(f"成功合并：{os.path.basename(file_path)}, 估计页数: {page_count}")
                    except Exception as e:
                        error_msg = f"处理文件 {os.path.basename(file_path)} 时出错：{str(e)}"
                        self.log(error_msg)
                        continue

                except Exception as e:
                    error_msg = f"处理文件 {os.path.basename(file_path)} 时出错：{str(e)}"
                    self.log(error_msg)
                    continue

            # 保存合并后的文档
            self.log("保存合并后的文档...")
            merged_doc.save(output_path)
            
            # 清理临时文件
            for temp_file in temp_files:
                try:
                    os.remove(temp_file)
                    self.log(f"清理临时文件：{os.path.basename(temp_file)}")
                except:
                    pass
                    
            return True

        except Exception as e:
            self.log(f"简单追加合并失败：{str(e)}")
            return False

    def algorithm_windows(self, files, final_docx):
        """Windows平台下的合并算法，增加关闭批注功能和页码记录"""
        from win32com import client
        word = None
        try:
            word = client.gencache.EnsureDispatch('Word.Application')
            word.Visible = False
            new_document = word.Documents.Add()
            current_page = 1
            self.file_page_map = {}  # 重置文件页码映射
            
            # 添加一个空白页用于目录
            new_document.Content.InsertAfter("\n")
            
            for i, fn in enumerate(files):
                self.log(f"正在合并文件：{os.path.basename(fn)}")
                fn = os.path.abspath(fn)
                
                # 获取当前页码
                current_page = new_document.ComputeStatistics(2)  # 2 = wdStatisticPages
                
                # 创建唯一书签名
                bookmark_name = f"bookmark_{i+1}"
                
                # 在当前位置添加书签
                end_position = new_document.Content.End - 1
                new_document.Bookmarks.Add(bookmark_name, new_document.Range(end_position, end_position))
                
                # 记录当前页码和书签
                self.file_page_map[fn] = {
                    'page': current_page,
                    'bookmark': bookmark_name
                }
                
                # 打开文档
                temp_document = word.Documents.Open(fn)
                
                # 关闭批注功能
                if temp_document.TrackRevisions:
                    temp_document.TrackRevisions = False
                
                # 接受所有修订
                temp_document.Revisions.AcceptAll()
                
                # 复制内容
                temp_document.Content.Copy()
                
                # 移动到文档末尾并粘贴
                end_range = new_document.Range(new_document.Content.End - 1, new_document.Content.End - 1)
                end_range.Paste()
                
                # 添加分页符（除了最后一个文档）
                if i < len(files) - 1:
                    end_range = new_document.Range(new_document.Content.End - 1, new_document.Content.End - 1)
                    end_range.InsertBreak(7)  # 7 = wdPageBreak
                
                # 关闭临时文档
                temp_document.Close(SaveChanges=False)
                self.log(f"成功合并：{os.path.basename(fn)}")
            
            # 保存合并后的文档
            self.log("保存合并后的文档...")
            new_document.SaveAs(final_docx)
            new_document.Close(SaveChanges=False)
            return True
        except Exception as e:
            error_msg = f"算法错误: {str(e)}"
            self.log(error_msg)
            messagebox.showerror("错误", error_msg)
            return False
        finally:
            # 确保在任何情况下都关闭Word
            if word:
                word.Quit()

    # 修改其他合并方法，添加关闭批注功能和页码记录
    def merge_with_word_api(self, output_path, doc_files):
        """使用 Word API 合并算法，增加关闭批注功能和页码记录"""
        word = None
        try:
            word = win32.gencache.EnsureDispatch("Word.Application")
            word.Visible = False
            merged_doc = word.Documents.Add()
            self.file_page_map = {}  # 重置文件页码映射
            
            # 添加一个空白页用于目录
            merged_doc.Content.InsertAfter("\n")
            
            for i, file_path in enumerate(doc_files):
                try:
                    self.log(f"正在合并文件：{os.path.basename(file_path)}")
                    
                    # 获取当前页码
                    current_page = merged_doc.ComputeStatistics(2)  # 2 = wdStatisticPages
                    
                    # 创建唯一书签名
                    bookmark_name = f"bookmark_{i+1}"
                    
                    # 在当前位置添加书签
                    end_position = merged_doc.Content.End - 1
                    merged_doc.Bookmarks.Add(bookmark_name, merged_doc.Range(end_position, end_position))
                    
                    # 记录当前页码和书签
                    self.file_page_map[file_path] = {
                        'page': current_page,
                        'bookmark': bookmark_name
                    }
                    
                    self.log(f"添加书签: {bookmark_name}, 页码: {current_page}, 文件: {os.path.basename(file_path)}")
                    
                    # 打开文档
                    try:
                        doc = word.Documents.Open(file_path)
                        
                        # 关闭批注功能
                        if doc.TrackRevisions:
                            doc.TrackRevisions = False
                        
                        # 接受所有修订
                        doc.Revisions.AcceptAll()
                        
                        # 复制内容
                        doc.Content.Copy()
                        
                        # 移动到文档末尾并粘贴
                        end_range = merged_doc.Range(merged_doc.Content.End - 1, merged_doc.Content.End - 1)
                        end_range.Paste()
                        
                        # 添加分页符（除了最后一个文档）
                        if i < len(doc_files) - 1:
                            end_range = merged_doc.Range(merged_doc.Content.End - 1, merged_doc.Content.End - 1)
                            end_range.InsertBreak(7)  # 7 = wdPageBreak
                        
                        # 关闭临时文档
                        doc.Close(SaveChanges=False)
                        self.log(f"成功合并：{os.path.basename(file_path)}")
                    except Exception as e:
                        error_msg = f"处理文件 {os.path.basename(file_path)} 时出错：{str(e)}"
                        self.log(error_msg)
                        continue
                        
                except Exception as e:
                    error_msg = f"处理文件 {os.path.basename(file_path)} 时出错：{str(e)}"
                    self.log(error_msg)
                    continue

            # 保存合并后的文档
            self.log("保存合并后的文档...")
            merged_doc.SaveAs(output_path)
            merged_doc.Close(SaveChanges=False)
            return True

        except Exception as e:
            self.log(f"Word API 合并失败：{str(e)}")
            return False
        finally:
            # 确保在任何情况下都关闭Word
            try:
                if word:
                    word.Quit()
            except:
                pass

    def merge_with_format(self, output_path, doc_files):
        """保留格式合并算法，使用python-docx和docxcompose库"""
        try:
            # 筛选出可以处理的.docx文件
            valid_files = []
            self.file_page_map = {}  # 重置文件页码映射
            current_page = 0
            
            # 预处理：将.doc文件转换为.docx
            temp_files = []
            for i, file_path in enumerate(doc_files):
                try:
                    self.log(f"正在处理文件：{os.path.basename(file_path)}")
                    
                    # 如果是.doc文件，先转换为.docx
                    if file_path.lower().endswith('.doc'):
                        self.log(f"转换.doc文件为.docx: {os.path.basename(file_path)}")
                        # 使用Word COM接口转换
                        word = win32.Dispatch("Word.Application")
                        word.Visible = False
                        try:
                            # 打开.doc文件
                            doc = word.Documents.Open(file_path)
                            # 创建临时.docx文件路径
                            temp_docx = os.path.join(os.path.dirname(output_path), f"temp_{i}_{os.path.basename(file_path)}x")
                            # 保存为.docx
                            doc.SaveAs(temp_docx, 16)  # 16 = wdFormatDocumentDefault (.docx)
                            doc.Close()
                            word.Quit()
                            # 添加到临时文件列表
                            temp_files.append(temp_docx)
                            # 使用转换后的文件
                            file_path = temp_docx
                        except Exception as e:
                            self.log(f"转换文件失败: {str(e)}")
                            word.Quit()
                            continue
                    
                    # 创建唯一书签名
                    bookmark_name = f"bookmark_{i+1}"
                    
                    # 记录当前页码和书签
                    self.file_page_map[file_path] = {
                        'page': current_page,
                        'bookmark': bookmark_name
                    }
                    
                    # 尝试打开文件验证其有效性
                    try:
                        doc = Document(file_path)
                        # 估算页数 (粗略计算，每页约2000个字符)
                        text_length = 0
                        for para in doc.paragraphs:
                            text_length += len(para.text)
                        # 更新当前页码 (粗略估计)
                        page_count = max(1, text_length // 2000)
                        current_page += page_count
                        
                        valid_files.append(file_path)
                        self.log(f"成功验证文件：{os.path.basename(file_path)}, 估计页数: {page_count}")
                    except Exception as e:
                        self.log(f"无法打开文件 {os.path.basename(file_path)}: {str(e)}")
                        continue
                    
                except Exception as e:
                    error_msg = f"处理文件 {os.path.basename(file_path)} 时出错：{str(e)}"
                    self.log(error_msg)
                    continue
            
            if not valid_files:
                self.log("没有有效的文件可以合并")
                return False
                
            # 使用docxcompose合并有效的文件
            self.log("开始合并有效的文件...")
            merged_doc = Document()
            composer = Composer(merged_doc)
            
            for file_path in valid_files:
                try:
                    self.log(f"合并文件：{os.path.basename(file_path)}")
                    doc = Document(file_path)
                    composer.append(doc)
                    self.log(f"成功合并：{os.path.basename(file_path)}")
                except Exception as e:
                    self.log(f"合并文件 {os.path.basename(file_path)} 时出错：{str(e)}")
            
            # 保存合并后的文档
            self.log("保存合并后的文档...")
            composer.save(output_path)
            
            # 清理临时文件
            for temp_file in temp_files:
                try:
                    os.remove(temp_file)
                except:
                    pass
                    
            return True
            
        except Exception as e:
            self.log(f"保留格式合并失败：{str(e)}")
            return False

    def merge_with_docxcompose(self, output_path, doc_files):
        """使用 docxcompose 合并算法，增加页码记录和书签支持"""
        try:
            # 筛选出可以处理的.docx文件
            valid_files = []
            self.file_page_map = {}  # 重置文件页码映射
            current_page = 0
            
            # 预处理：将.doc文件转换为.docx
            temp_files = []
            for i, file_path in enumerate(doc_files):
                try:
                    self.log(f"正在处理文件：{os.path.basename(file_path)}")
                    file_path = os.path.abspath(file_path)  # 确保使用绝对路径
                    
                    # 如果是.doc文件，先转换为.docx
                    if file_path.lower().endswith('.doc'):
                        self.log(f"转换.doc文件为.docx: {os.path.basename(file_path)}")
                        # 使用Word COM接口转换
                        word = None
                        try:
                            word = win32.Dispatch("Word.Application")
                            word.Visible = False
                            
                            # 打开.doc文件
                            doc = word.Documents.Open(file_path)
                            # 创建临时.docx文件路径
                            temp_docx = os.path.join(os.path.dirname(output_path), f"temp_{i}_{os.path.basename(file_path)}x")
                            # 保存为.docx
                            doc.SaveAs(temp_docx, 16)  # 16 = wdFormatDocumentDefault (.docx)
                            doc.Close(SaveChanges=False)
                            
                            # 添加到临时文件列表
                            temp_files.append(temp_docx)
                            # 使用转换后的文件
                            file_path = temp_docx
                            self.log(f"成功转换文件：{os.path.basename(file_path)}")
                        except Exception as e:
                            self.log(f"转换文件失败: {str(e)}")
                            continue
                        finally:
                            if word:
                                try:
                                    word.Quit()
                                except:
                                    pass
                    
                    # 创建唯一书签名
                    bookmark_name = f"bookmark_{i+1}"
                    
                    # 尝试打开文件验证其有效性
                    try:
                        doc = Document(file_path)
                        
                        # 估算页数 (粗略计算，每页约2000个字符)
                        text_length = 0
                        for para in doc.paragraphs:
                            text_length += len(para.text)
                        # 更新当前页码 (粗略估计)
                        page_count = max(1, text_length // 2000)
                        
                        # 记录当前页码和书签
                        self.file_page_map[file_path] = {
                            'page': current_page,
                            'bookmark': bookmark_name
                        }
                        
                        current_page += page_count
                        valid_files.append(file_path)
                        self.log(f"成功验证文件：{os.path.basename(file_path)}, 估计页数: {page_count}")
                    except Exception as e:
                        self.log(f"无法打开文件 {os.path.basename(file_path)}: {str(e)}")
                        continue
                    
                except Exception as e:
                    error_msg = f"处理文件 {os.path.basename(file_path)} 时出错：{str(e)}"
                    self.log(error_msg)
                    continue
            
            if not valid_files:
                self.log("没有有效的文件可以合并")
                return False
                
            # 使用docxcompose合并有效的文件
            self.log("开始合并有效的文件...")
            merged_doc = Document()
            composer = Composer(merged_doc)
            
            for file_path in valid_files:
                try:
                    self.log(f"合并文件：{os.path.basename(file_path)}")
                    doc = Document(file_path)
                    composer.append(doc)
                    self.log(f"成功合并：{os.path.basename(file_path)}")
                except Exception as e:
                    self.log(f"合并文件 {os.path.basename(file_path)} 时出错：{str(e)}")
                    continue
            
            # 保存合并后的文档
            self.log("保存合并后的文档...")
            composer.save(output_path)
            
            # 清理临时文件
            for temp_file in temp_files:
                try:
                    os.remove(temp_file)
                    self.log(f"清理临时文件：{os.path.basename(temp_file)}")
                except:
                    pass
                    
            return True
            
        except Exception as e:
            self.log(f"docxcompose 合并失败：{str(e)}")
            return False

if __name__ == "__main__":
    app = WordMergerApp()
    app.mainloop()